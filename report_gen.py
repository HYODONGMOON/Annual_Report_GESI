import os
import re
import unicodedata
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

EXCEL_PATH   = 'Annual Report Database.xlsx'
TOC_IMAGE    = 'toc_background.png'
PARTNERS_DIR = 'partners'

# 발행물 접두어 순서 및 표시 명칭
PUBLICATION_CATEGORIES = [
    ('[연구보고서]', '연구보고서'),
    ('[이슈리포트]', '이슈리포트'),
    ('[논문]',      '논문'),
    ('[기타]',      '기타'),
]


def _normalize(text: str) -> str:
    """파일명·기관명 비교용 정규화 (소문자, 공백·특수문자 제거)"""
    text = unicodedata.normalize('NFC', text)
    return re.sub(r'[\s\-_.,()·/]+', '', text).lower()


class GesiFullReportGenerator:
    def __init__(self, excel_path=EXCEL_PATH):
        self.doc = Document()
        self.xl  = pd.ExcelFile(excel_path)

        style = self.doc.styles['Normal']
        style.font.name = '맑은 고딕'
        style.font.size = Pt(11)

    def _read_sheet(self, sheet_name, **kwargs):
        return pd.read_excel(self.xl, sheet_name=sheet_name, **kwargs)

    # ------------------------------------------------------------------ #
    #  헬퍼: Pillar 번호 추출                                               #
    # ------------------------------------------------------------------ #
    @staticmethod
    def _pillar_key(field_str: str) -> str:
        """'1. 시스템 ...' → '1'"""
        m = re.match(r'(\d+)', str(field_str).strip())
        return m.group(1) if m else ''

    # ------------------------------------------------------------------ #
    #  1. 표지                                                              #
    # ------------------------------------------------------------------ #
    def add_title_page(self):
        for _ in range(5):
            self.doc.add_paragraph()

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("2025 GESI ANNUAL REPORT")
        run.bold = True
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(34, 139, 34)

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run("탄소중립을 향한 데이터와 현장의 기록")
        run2.font.size = Pt(16)

        self.doc.add_page_break()

    # ------------------------------------------------------------------ #
    #  2. 목차 (이미지 포함)                                                #
    # ------------------------------------------------------------------ #
    def add_toc_page(self, image_path=TOC_IMAGE):
        if image_path and os.path.exists(image_path):
            p_img = self.doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(image_path, width=Inches(6.2))
        else:
            self.doc.add_paragraph()

        self.doc.add_paragraph()

        toc_title = self.doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = toc_title.add_run("목  차  /  CONTENTS")
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(34, 139, 34)

        self.doc.add_paragraph()

        toc_items = [
            ("01", "기관 소개",              None),
            ("02", "연구 연혁 (2015–2025)",  None),
            ("03", "주요 활동 및 외부 소통",  None),
            ("04", "2025년 주요 연구 상세",   "2025"),   # Pillar 소항목 자동 삽입
            ("05", "협력기관",               None),
            ("06", "발행물",                 None),
        ]

        for num, item_title, sub_sheet in toc_items:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.8)
            r_num = p.add_run(f"{num}.  ")
            r_num.bold = True
            r_num.font.size = Pt(12)
            r_num.font.color.rgb = RGBColor(34, 139, 34)
            r_body = p.add_run(item_title)
            r_body.font.size = Pt(12)

            if sub_sheet:
                try:
                    df_sub = self._read_sheet(sub_sheet, index_col=0)
                    for col in df_sub.columns:
                        p_sub = self.doc.add_paragraph()
                        p_sub.paragraph_format.left_indent = Inches(1.4)
                        r_sub = p_sub.add_run(f"∙  {col}")
                        r_sub.font.size = Pt(10)
                        r_sub.font.color.rgb = RGBColor(90, 90, 90)
                except Exception:
                    pass

        self.doc.add_page_break()

    # ------------------------------------------------------------------ #
    #  3. 기관 소개 (Institute_Info)                                        #
    # ------------------------------------------------------------------ #
    def add_institute_intro(self):
        self.doc.add_heading('01. 기관 소개', level=1)
        try:
            df = self._read_sheet('Institute_Info')
            row_2025 = df[df.iloc[:, 0] == 2025]
            row = row_2025.iloc[0] if not row_2025.empty else df.iloc[0]

            col_overview = df.columns[1]
            col_people   = df.columns[2]
            col_contact  = df.columns[3]

            if pd.notna(row[col_overview]):
                p = self.doc.add_paragraph(str(row[col_overview]))
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            self.doc.add_paragraph()

            tbl = self.doc.add_table(rows=2, cols=2)
            tbl.style = 'Light Grid Accent 3'
            tbl.rows[0].cells[0].text = '기관 인력'
            tbl.rows[0].cells[1].text = str(row[col_people])  if pd.notna(row[col_people])  else '-'
            tbl.rows[1].cells[0].text = '주요 연락처'
            tbl.rows[1].cells[1].text = str(row[col_contact]) if pd.notna(row[col_contact]) else '-'
        except Exception as e:
            self.doc.add_paragraph(f"기관 소개 데이터를 불러올 수 없습니다. ({e})")

        self.doc.add_paragraph()
        self.doc.add_page_break()

    # ------------------------------------------------------------------ #
    #  4. 연구 연혁 (Research_History + 타임라인 인포그래픽)                #
    # ------------------------------------------------------------------ #
    def _create_timeline_infographic(self) -> str:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches

        for fname in ('Malgun Gothic', 'NanumGothic', 'AppleGothic', 'sans-serif'):
            try:
                matplotlib.rc('font', family=fname)
                break
            except Exception:
                continue
        matplotlib.rc('axes', unicode_minus=False)

        df_hist   = self._read_sheet('Research_History')
        df_master = self._read_sheet('Master_Research')

        year_col    = df_hist.columns[0]
        pillar_cols = df_hist.columns[1:]
        df_hist = df_hist.copy()
        df_hist[year_col] = df_hist[year_col].ffill()

        yp_count: dict = {}
        for _, row in df_hist.iterrows():
            yr = int(row[year_col])
            for p_idx, p_col in enumerate(pillar_cols):
                if pd.notna(row[p_col]):
                    yp_count[(yr, p_idx)] = yp_count.get((yr, p_idx), 0) + 1

        col_id     = df_master.columns[0]
        col_parent = df_master.columns[1]
        id_to_pos: dict = {}
        for _, row in df_master.iterrows():
            rid = str(row[col_id]).strip()
            parts = rid.split('-')
            if len(parts) == 4:
                try:
                    id_to_pos[rid] = (int(parts[1]), int(parts[2]) - 1)
                except ValueError:
                    pass

        connections = []
        for _, row in df_master.iterrows():
            rid    = str(row[col_id]).strip()
            parent = row[col_parent]
            if pd.notna(parent):
                pid = str(parent).strip()
                if rid in id_to_pos and pid in id_to_pos:
                    connections.append((id_to_pos[pid], id_to_pos[rid]))

        fig, ax = plt.subplots(figsize=(16, 6.5))
        fig.patch.set_facecolor('#f0f4f0')
        ax.set_facecolor('#f0f4f0')

        PILLAR_COLORS = ['#2e7d32', '#1565c0', '#bf360c']
        PILLAR_Y      = [4.0, 2.0, 0.0]
        PILLAR_LABELS = [
            'Pillar 1\n시스템 전환\n& 섹터커플링',
            'Pillar 2\n지역 에너지 전환\n& 분산에너지',
            'Pillar 3\n탈탄소 정책\n& 사회적 가치',
        ]
        YEARS = list(range(2015, 2026))

        for p_idx in range(3):
            y = PILLAR_Y[p_idx]
            rect = mpatches.FancyBboxPatch(
                (2014.7, y - 0.72), 10.8, 1.44,
                boxstyle='round,pad=0.05', linewidth=0,
                facecolor=PILLAR_COLORS[p_idx], alpha=0.10
            )
            ax.add_patch(rect)
            ax.text(2014.35, y, PILLAR_LABELS[p_idx],
                    ha='right', va='center', fontsize=8,
                    color=PILLAR_COLORS[p_idx], fontweight='bold', linespacing=1.5)

        for yr in YEARS:
            ax.axvline(x=yr, color='#bbbbbb', linewidth=0.6,
                       linestyle=':', alpha=0.9, zorder=1)
            ax.text(yr, -1.1, str(yr), ha='center', va='top',
                    fontsize=7.5, color='#666666')

        for (yr, p_idx), cnt in yp_count.items():
            y_base = PILLAR_Y[p_idx]
            size   = 60 + cnt * 20
            ax.scatter(yr, y_base, s=min(size, 220),
                       color=PILLAR_COLORS[p_idx], alpha=0.75,
                       zorder=5, edgecolors='white', linewidth=1.5)
            if cnt > 1:
                ax.text(yr, y_base, str(cnt), ha='center', va='center',
                        fontsize=6.5, color='white', fontweight='bold', zorder=6)

        for (from_yr, from_p), (to_yr, to_p) in connections:
            fy, ty = PILLAR_Y[from_p], PILLAR_Y[to_p]
            rad = 0.35 if from_p == to_p else -0.45
            ax.annotate('',
                xy=(to_yr, ty), xytext=(from_yr, fy),
                arrowprops=dict(
                    arrowstyle='->', color='#e65100', lw=2.2,
                    connectionstyle=f'arc3,rad={rad}',
                    shrinkA=9, shrinkB=9
                ), zorder=7)

        for p_idx in range(3):
            y = PILLAR_Y[p_idx]
            ax.annotate('', xy=(2025.6, y), xytext=(2014.8, y),
                        arrowprops=dict(arrowstyle='->',
                                        color=PILLAR_COLORS[p_idx],
                                        lw=1.2, alpha=0.35))

        ax.set_title('GESI 연구 연혁 타임라인  (2015–2025)',
                     fontsize=14, fontweight='bold', color='#1a3d2b', pad=14)

        legend_handles = [
            mpatches.Patch(color=PILLAR_COLORS[0], alpha=0.8,
                           label='Pillar 1: 시스템 전환 & 섹터커플링'),
            mpatches.Patch(color=PILLAR_COLORS[1], alpha=0.8,
                           label='Pillar 2: 지역 에너지 전환 & 분산에너지'),
            mpatches.Patch(color=PILLAR_COLORS[2], alpha=0.8,
                           label='Pillar 3: 탈탄소 정책 & 사회적 가치'),
            mpatches.Patch(color='#e65100', alpha=0.8,
                           label='연구 연계 / 확장 (연구 분야 ID 기반)'),
        ]
        ax.legend(handles=legend_handles, loc='upper right',
                  fontsize=8.5, framealpha=0.9, edgecolor='#cccccc')

        ax.text(2015, -1.75,
                '※ 버블 크기 = 해당 연도·분야 연구 수  |  숫자 = 연구 건수',
                fontsize=7.5, color='#888888')

        ax.set_xlim(2013.8, 2026.2)
        ax.set_ylim(-2.1, 5.5)
        ax.axis('off')
        plt.tight_layout(pad=1.5)

        out_path = 'research_timeline.png'
        plt.savefig(out_path, dpi=150, bbox_inches='tight',
                    facecolor=fig.get_facecolor())
        plt.close()
        return out_path

    def add_research_history(self):
        self.doc.add_heading('02. 연구 연혁 (2015–2025)', level=1)
        self.doc.add_paragraph(
            "GESI가 걸어온 지난 10년은 한국 에너지 전환의 역사와 궤를 같이 합니다. "
            "아래 타임라인은 세 가지 핵심 연구 축이 어떻게 심화·확장되어 왔는지를 보여줍니다."
        )
        self.doc.add_paragraph()

        try:
            infographic_path = self._create_timeline_infographic()
            p_img = self.doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(infographic_path, width=Inches(6.4))
        except Exception as e:
            self.doc.add_paragraph(f"[인포그래픽 생성 오류: {e}]")

        self.doc.add_paragraph()

        try:
            df = self._read_sheet('Research_History')
            year_col    = df.columns[0]
            pillar_cols = df.columns[1:]
            df = df.copy()
            df[year_col] = df[year_col].ffill()

            tbl = self.doc.add_table(rows=1, cols=1 + len(pillar_cols))
            tbl.style = 'Light Grid Accent 3'
            hdr = tbl.rows[0].cells
            hdr[0].text = '연도'
            for i, col in enumerate(pillar_cols):
                hdr[i + 1].text = str(col)

            for _, row in df.iterrows():
                cells = tbl.add_row().cells
                yr_val = row[year_col]
                cells[0].text = str(int(yr_val)) if pd.notna(yr_val) else ''
                for i, col in enumerate(pillar_cols):
                    cells[i + 1].text = str(row[col]) if pd.notna(row[col]) else ''
        except Exception as e:
            self.doc.add_paragraph(f"연혁 데이터를 불러올 수 없습니다. ({e})")

        self.doc.add_page_break()

    # ------------------------------------------------------------------ #
    #  5. 주요 활동 (Activities_&_News)  ← 구 04                           #
    # ------------------------------------------------------------------ #
    def add_activities(self):
        self.doc.add_heading('03. 주요 활동 및 외부 소통', level=1)
        self.doc.add_paragraph(
            "GESI는 연구를 넘어 시민사회 및 정책 입안자들과 끊임없이 소통하고 있습니다."
        )
        try:
            df = self._read_sheet('Activities_&_News')
            col_date    = df.columns[1]
            col_type    = df.columns[3]
            col_name    = df.columns[4]
            col_content = df.columns[6]

            for _, row in df.iterrows():
                date_str = str(row[col_date]) if pd.notna(row[col_date]) else ''
                act_name = str(row[col_name]) if pd.notna(row[col_name]) else ''
                act_type = str(row[col_type]) if pd.notna(row[col_type]) else ''
                p = self.doc.add_paragraph()
                p.add_run(f"• [{date_str}] {act_name}  ({act_type})").bold = True
                if pd.notna(row[col_content]):
                    self.doc.add_paragraph(f"  : {row[col_content]}")
        except Exception as e:
            self.doc.add_paragraph(f"활동 데이터를 불러올 수 없습니다. ({e})")

        self.doc.add_page_break()

    # ------------------------------------------------------------------ #
    #  6. 2025년 주요 연구 상세 — Pillar별 (구 05 + 구 03 병합)            #
    #     각 Pillar 페이지: 상세 서술(2025 시트) + 연구 목록(Master_Research)#
    # ------------------------------------------------------------------ #
    def add_2025_pillar_pages(self):
        self.doc.add_heading('04. 2025년 주요 연구 상세', level=1)

        try:
            df_detail = self._read_sheet('2025', index_col=0)
            df_master = self._read_sheet('Master_Research')

            col_field   = df_master.columns[3]   # 연구 분야
            col_title   = df_master.columns[4]   # 연구명
            col_summary = df_master.columns[5]   # 다시 말해

            # Pillar별 연구 목록 미리 구성
            pillar_research: dict = {}
            for _, row in df_master.iterrows():
                key = self._pillar_key(str(row[col_field]))
                if key:
                    pillar_research.setdefault(key, []).append(row)

            for pillar_name in df_detail.columns:
                # ── 페이지 상단: Pillar 헤딩 ──────────────────────────────
                heading = self.doc.add_heading('', level=2)
                run = heading.add_run(str(pillar_name))
                run.font.color.rgb = RGBColor(34, 139, 34)

                # ── 2025 시트 상세 서술 ───────────────────────────────────
                for row_label, cell_value in df_detail[pillar_name].items():
                    if pd.isna(cell_value):
                        continue
                    lp = self.doc.add_paragraph()
                    lp.add_run(str(row_label).strip()).bold = True
                    cp = self.doc.add_paragraph(str(cell_value).strip().strip('"'))
                    cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    self.doc.add_paragraph()

                # ── 해당 Pillar 연구 목록 (Master_Research) ───────────────
                p_key = self._pillar_key(pillar_name)
                if p_key and p_key in pillar_research:
                    sep = self.doc.add_paragraph()
                    sep.add_run("▌ 관련 연구 목록").bold = True

                    tbl = self.doc.add_table(rows=1, cols=2)
                    tbl.style = 'Light Grid Accent 3'
                    tbl.rows[0].cells[0].text = '연구명'
                    tbl.rows[0].cells[1].text = '핵심 요약'

                    for row in pillar_research[p_key]:
                        cells = tbl.add_row().cells
                        cells[0].text = str(row[col_title]) if pd.notna(row[col_title]) else '-'
                        cells[1].text = str(row[col_summary]) if pd.notna(row[col_summary]) else '-'

                self.doc.add_page_break()

        except Exception as e:
            self.doc.add_paragraph(f"2025 Pillar 데이터를 불러올 수 없습니다. ({e})")

    # ------------------------------------------------------------------ #
    #  7. 협력기관 (주요 발주처 → partners/ 폴더 심볼 로드)                 #
    # ------------------------------------------------------------------ #
    def add_partners_page(self, partners_dir=PARTNERS_DIR):
        """
        Master_Research 의 '주요 발주처' 컬럼에서 기관명을 추출한 뒤,
        partners/ 폴더에서 기관명과 일치하는 이미지 파일을 찾아 삽입합니다.

        ■ 이미지 파일 네이밍 규칙 (partners/ 폴더):
          - 파일명에 기관명 일부가 포함되면 자동 매핑됩니다.
          - 예) '환경부.png', 'TARA_logo.png', '산업통상자원부.jpg' 등
          - 지원 확장자: .png, .jpg, .jpeg, .bmp, .gif
        """
        self.doc.add_heading('05. 협력기관', level=1)
        self.doc.add_paragraph(
            "본 연구소는 아래 기관들과의 협력을 통해 연구를 수행하였습니다."
        )
        self.doc.add_paragraph()

        try:
            df = self._read_sheet('Master_Research')
            col_partner = df.columns[6]   # 주요 발주처

            # 기관명 추출 및 중복 제거 (쉼표/슬래시 구분자 처리)
            orgs: list[str] = []
            seen: set = set()
            for val in df[col_partner].dropna():
                for org in re.split(r'[,/\n]+', str(val)):
                    org = org.strip()
                    if org and org not in seen:
                        orgs.append(org)
                        seen.add(org)

            if not orgs:
                self.doc.add_paragraph("(협력기관 정보가 없습니다.)")
                self.doc.add_page_break()
                return

            # partners/ 폴더의 이미지 파일 목록
            img_exts = {'.png', '.jpg', '.jpeg', '.bmp', '.gif'}
            partner_files: dict[str, str] = {}   # normalized_name → full_path
            if os.path.isdir(partners_dir):
                for fname in os.listdir(partners_dir):
                    stem, ext = os.path.splitext(fname)
                    if ext.lower() in img_exts:
                        partner_files[_normalize(stem)] = os.path.join(partners_dir, fname)

            def _find_image(org_name: str) -> str | None:
                """기관명과 파일명을 정규화 비교하여 이미지 경로 반환"""
                norm_org = _normalize(org_name)
                # 완전 일치 우선
                if norm_org in partner_files:
                    return partner_files[norm_org]
                # 부분 포함 검색
                for norm_file, path in partner_files.items():
                    if norm_org in norm_file or norm_file in norm_org:
                        return path
                return None

            # 3열 그리드 레이아웃 (이미지 + 기관명)
            COLS = 3
            rows_needed = (len(orgs) + COLS - 1) // COLS
            tbl = self.doc.add_table(rows=rows_needed * 2, cols=COLS)
            tbl.style = 'Table Grid'

            for idx, org in enumerate(orgs):
                row_img  = (idx // COLS) * 2
                row_text = row_img + 1
                col_idx  = idx % COLS

                img_cell  = tbl.rows[row_img].cells[col_idx]
                text_cell = tbl.rows[row_text].cells[col_idx]

                img_path = _find_image(org)
                if img_path:
                    p_img = img_cell.paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        p_img.add_run().add_picture(img_path, width=Inches(1.3))
                    except Exception:
                        img_cell.text = '[이미지 로드 실패]'
                else:
                    # 이미지 없으면 기관명만 (회색 박스 효과)
                    p_ph = img_cell.paragraphs[0]
                    p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_ph = p_ph.add_run(org)
                    r_ph.font.color.rgb = RGBColor(150, 150, 150)
                    r_ph.font.size = Pt(9)

                p_name = text_cell.paragraphs[0]
                p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_name = p_name.add_run(org)
                r_name.bold = True
                r_name.font.size = Pt(9)

        except Exception as e:
            self.doc.add_paragraph(f"협력기관 데이터를 불러올 수 없습니다. ({e})")

        self.doc.add_page_break()

    # ------------------------------------------------------------------ #
    #  8. 발행물 ([연구보고서] / [이슈리포트] / [논문] / [기타])            #
    # ------------------------------------------------------------------ #
    def add_publications_page(self):
        """
        Master_Research 의 '주요 성과물' 컬럼을 접두어 기준으로 분류합니다.

        ■ 입력 형식 (엑셀 셀):
          [연구보고서] 보고서 제목
          [이슈리포트] 이슈리포트 제목
          [논문] 논문 제목
          [기타] 기타 성과물

          → 한 셀에 여러 항목이 있을 경우 줄바꿈(\n)으로 구분해 주세요.
        """
        self.doc.add_heading('06. 발행물', level=1)
        self.doc.add_paragraph(
            "2025년 GESI의 연구 성과물을 유형별로 정리하였습니다."
        )
        self.doc.add_paragraph()

        # 카테고리별 버킷
        buckets: dict[str, list[str]] = {tag: [] for tag, _ in PUBLICATION_CATEGORIES}
        buckets['[기타]'] = buckets.get('[기타]', [])

        try:
            df = self._read_sheet('Master_Research')
            col_output = df.columns[7]   # 주요 성과물
            col_title  = df.columns[4]   # 연구명 (출처 표기용)

            for _, row in df.iterrows():
                cell = row[col_output]
                if pd.isna(cell):
                    continue
                research_name = str(row[col_title]) if pd.notna(row[col_title]) else ''

                # 줄바꿈으로 분리된 항목 처리
                for item in str(cell).splitlines():
                    item = item.strip()
                    if not item:
                        continue
                    matched = False
                    for tag, _ in PUBLICATION_CATEGORIES:
                        if item.startswith(tag):
                            pub_title = item[len(tag):].strip()
                            entry = f"{pub_title}  ← {research_name}" if research_name else pub_title
                            buckets[tag].append(entry)
                            matched = True
                            break
                    if not matched and item:
                        # 접두어 없는 항목은 [기타]로
                        entry = f"{item}  ← {research_name}" if research_name else item
                        buckets['[기타]'].append(entry)

        except Exception as e:
            self.doc.add_paragraph(f"발행물 데이터를 불러올 수 없습니다. ({e})")
            self.doc.add_page_break()
            return

        # 카테고리별 출력
        has_any = False
        for tag, display_name in PUBLICATION_CATEGORIES:
            items = buckets.get(tag, [])
            if not items:
                continue
            has_any = True

            # 카테고리 소제목
            h = self.doc.add_heading('', level=2)
            r = h.add_run(f"■ {display_name}")
            r.font.color.rgb = RGBColor(34, 139, 34)

            tbl = self.doc.add_table(rows=1, cols=2)
            tbl.style = 'Light Grid Accent 3'
            tbl.rows[0].cells[0].text = 'No.'
            tbl.rows[0].cells[1].text = '제목 / 출처 연구'

            for i, entry in enumerate(items, start=1):
                cells = tbl.add_row().cells
                cells[0].text = str(i)
                cells[1].text = entry

            self.doc.add_paragraph()

        if not has_any:
            self.doc.add_paragraph(
                "(아직 등록된 발행물이 없습니다. "
                "엑셀 'Master_Research' 시트의 '주요 성과물' 열에 "
                "[연구보고서] / [이슈리포트] / [논문] / [기타] 형식으로 입력해 주세요.)"
            )

        self.doc.add_page_break()

    # ------------------------------------------------------------------ #
    #  저장                                                                #
    # ------------------------------------------------------------------ #
    def save_report(self, filename):
        self.doc.save(filename)
        print(f"보고서 생성 완료: {filename}")


# ─────────────────────────── 실행부 ────────────────────────────────────── #
if __name__ == "__main__":
    gen = GesiFullReportGenerator(EXCEL_PATH)
    gen.add_title_page()
    gen.add_toc_page(TOC_IMAGE)
    gen.add_institute_intro()
    gen.add_research_history()
    gen.add_activities()
    gen.add_2025_pillar_pages()
    gen.add_partners_page(PARTNERS_DIR)
    gen.add_publications_page()
    gen.save_report("GESI_2025_Annual_Report_Final.docx")
